"""Universal document text extraction for bill ingestion.

Routes any uploaded file to the right reader and returns plain text:
  • PDF        → text layer (pdfplumber/pypdf); OCR fallback if scanned
  • Images     → OCR (png/jpg/jpeg/tiff/bmp/webp/gif/heic-as-image)
  • Excel      → .xlsx/.xlsm (openpyxl), .xls (xlrd)
  • Word       → .docx (python-docx)
  • CSV/TSV/TXT→ decoded text
  • HTML       → tag-stripped text
  • anything else → best-effort UTF-8 decode

OCR uses RapidOCR (bundled ONNX models — no system binary, works offline).
"""
from __future__ import annotations

import io
import re
from functools import lru_cache

IMAGE_EXT = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".gif"}


@lru_cache(maxsize=1)
def _ocr_engine():
    from rapidocr_onnxruntime import RapidOCR
    return RapidOCR()


def _ocr_array(arr) -> str:
    """OCR an image and reconstruct visual rows.

    RapidOCR returns one entry per detected text box, so a label and its amount
    ('Taxable Value:' / '1,20,000.00') arrive as separate items. We regroup
    boxes that share a vertical band into a single line, left-to-right — so the
    downstream 'label ... value on one line' extractor works on scans too.
    """
    result, _ = _ocr_engine()(arr)
    if not result:
        return ""

    boxes = []
    for item in result:  # [box(4 pts), text, score]
        box, text = item[0], item[1]
        ys = [p[1] for p in box]
        xs = [p[0] for p in box]
        boxes.append({
            "yc": sum(ys) / len(ys), "xl": min(xs),
            "h": max(ys) - min(ys), "text": text,
        })
    heights = sorted(b["h"] for b in boxes)
    thr = max(heights[len(heights) // 2] * 0.6, 6)  # half a line-height band

    boxes.sort(key=lambda b: b["yc"])
    rows: list[dict] = []
    for b in boxes:
        row = next((r for r in rows if abs(b["yc"] - r["yc"]) <= thr), None)
        if row is None:
            rows.append({"yc": b["yc"], "items": [b]})
        else:
            row["items"].append(b)
            row["yc"] = sum(i["yc"] for i in row["items"]) / len(row["items"])
    rows.sort(key=lambda r: r["yc"])
    return "\n".join(
        " ".join(i["text"] for i in sorted(r["items"], key=lambda x: x["xl"]))
        for r in rows
    )


def _ocr_image_bytes(content: bytes) -> str:
    from PIL import Image
    import numpy as np
    img = Image.open(io.BytesIO(content)).convert("RGB")
    return _ocr_array(np.array(img))


def _pdf_text_layer(content: bytes) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            text = "\n".join((p.extract_text(x_tolerance=1.5) or "") for p in pdf.pages)
        if text.strip():
            return text
    except Exception:
        pass
    try:
        from pypdf import PdfReader
        text = "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(content)).pages)
        if text.strip():
            return text
    except Exception:
        pass
    return ""


def _pdf_ocr(content: bytes) -> str:
    """Rasterize each PDF page (pypdfium2, bundled pdfium) and OCR it."""
    import pypdfium2 as pdfium
    import numpy as np
    pdf = pdfium.PdfDocument(content)
    out = []
    try:
        for i in range(len(pdf)):
            page = pdf[i]
            pil = page.render(scale=2).to_pil().convert("RGB")  # ~144 dpi
            out.append(_ocr_array(np.array(pil)))
    finally:
        pdf.close()
    return "\n".join(out)


def _xlsx_text(content: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" ".join(cells))
    return "\n".join(lines)


def _xls_text(content: bytes) -> str:
    import xlrd
    book = xlrd.open_workbook(file_contents=content)
    lines = []
    for sh in book.sheets():
        for r in range(sh.nrows):
            cells = [str(c.value) for c in sh.row(r) if c.value not in (None, "")]
            if cells:
                lines.append(" ".join(cells))
    return "\n".join(lines)


def _docx_text(content: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(content))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            parts.append(" ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _html_text(content: bytes) -> str:
    txt = content.decode("utf-8", "ignore")
    txt = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", txt, flags=re.I | re.S)
    txt = re.sub(r"<[^>]+>", " ", txt)
    return re.sub(r"[ \t]+", " ", txt)


def extract(filename: str, content: bytes) -> tuple[str, str, bool]:
    """Return (text, detected_mode, ocr_used). Never raises for unknown types."""
    name = (filename or "").lower()
    ext = "." + name.rsplit(".", 1)[-1] if "." in name else ""
    ocr_used = False

    if ext == ".pdf":
        text = _pdf_text_layer(content)
        mode = "pdf"
        if not text.strip():
            text = _pdf_ocr(content)  # scanned PDF
            ocr_used = True
    elif ext in IMAGE_EXT:
        text, mode, ocr_used = _ocr_image_bytes(content), "image", True
    elif ext in (".xlsx", ".xlsm"):
        text, mode = _xlsx_text(content), "excel"
    elif ext == ".xls":
        text, mode = _xls_text(content), "excel"
    elif ext == ".docx":
        text, mode = _docx_text(content), "word"
    elif ext in (".csv", ".tsv"):
        text, mode = content.decode("utf-8", "ignore"), "csv"
    elif ext in (".htm", ".html"):
        text, mode = _html_text(content), "manual"
    else:
        text, mode = content.decode("utf-8", "ignore"), "manual"

    return text, mode, ocr_used
